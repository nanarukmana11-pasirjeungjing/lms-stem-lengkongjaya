# -*- coding: utf-8 -*-
"""
LMS STEM — Gugus Lengkongjaya
Titik masuk utama aplikasi Streamlit.

Struktur proyek (siap deploy — dokumen sumber dibundel di dalam folder app):
    lms_stem_app/
        app.py              <- entry point (halaman ini)
        data/
            content.py      <- semua teks & data konten, dipisah dari tampilan
        assets/
            style.css        <- styling mobile-first
        docs/                <- seluruh dokumen sumber (.docx) dibundel di sini
        requirements.txt

Catatan lokasi berkas sumber:
    Dokumen "Daftar_Proyek_STEM_Kelas*.docx", "Modul_Proyek_STEM_*.docx",
    "Panduan_Guru_Pembuatan_Produk_*.docx", dan "Program_Supervisi_Kepala_
    Sekolah_STEM.docx" dibaca dari folder docs/ di dalam folder aplikasi ini
    (lihat DOCS_DIR) setiap kali halaman terkait dibuka. Untuk memperbarui
    kontennya di server online, ganti berkas di folder docs/ lalu commit &
    push ulang ke GitHub — Streamlit Community Cloud akan otomatis redeploy.

Alur akses:
    1. Login sederhana (pilih peran + nama, tanpa kata sandi) -> menentukan
       menu apa saja yang tampil. Portal Supervisi hanya untuk peran
       "Kepala Sekolah".
    2. Menu navigasi radio horizontal di bagian atas untuk berpindah antar
       Beranda, Modul Kelas, dan Portal Supervisi.
    3. Portal Supervisi berisi 4 sub-halaman (Lampiran 1-4) sesuai struktur
       dokumen Program Supervisi Kepala Sekolah.
"""

import re
import streamlit as st
from pathlib import Path

from data.content import (
    OVERVIEW,
    EVENT,
    AGENDA,
    PRINSIP_WARNA,
    PENGALAMAN_WARNA,
    CLASS_FILES,
    KELOMPOK_KELAS,
    SUPERVISI_FILE,
    NAV_PAGES,
    ROLES,
    LAMPIRAN_TABS,
    PROJECT_DOCS,
    STORYBOOK_LINKS,
)

try:
    import docx  # python-docx — untuk membaca tabel/paragraf dari file .docx
except ImportError:
    docx = None

# ----------------------------------------------------------------------------
# Konfigurasi halaman (mobile-first: layout centered, sidebar collapsed)
# ----------------------------------------------------------------------------
st.set_page_config(
    page_title="LMS STEM — Gugus Lengkongjaya",
    page_icon="🧪",
    layout="centered",
    initial_sidebar_state="collapsed",
)

APP_DIR = Path(__file__).resolve().parent
DOCS_DIR = APP_DIR / "docs"  # dokumen KKG STEM dibundel di dalam folder aplikasi (siap deploy)


def load_css():
    css_path = APP_DIR / "assets" / "style.css"
    if css_path.exists():
        st.markdown(f"<style>{css_path.read_text(encoding='utf-8')}</style>", unsafe_allow_html=True)


def init_state():
    if "auth" not in st.session_state:
        st.session_state.auth = None  # {"role": ..., "name": ..., "kelas": int|None}
    if "page" not in st.session_state:
        st.session_state.page = "landing"
    if "kelas_terpilih" not in st.session_state:
        st.session_state.kelas_terpilih = 1
    if "lampiran_tab" not in st.session_state:
        st.session_state.lampiran_tab = "lampiran1"


# ----------------------------------------------------------------------------
# Halaman: Login
# ----------------------------------------------------------------------------
def render_login():
    st.markdown(
        """
        <div class="hero-banner">
            <div class="hero-eyebrow">🧪 LMS STEM • Gugus Lengkongjaya</div>
            <div class="hero-title">Masuk ke Aplikasi</div>
            <div class="hero-subtitle">Pilih peran Anda untuk melanjutkan</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="card"><p>Login sederhana — cukup pilih peran dan isi nama, tanpa kata sandi. '
        'Peran menentukan menu yang tersedia (Portal Supervisi khusus Kepala Sekolah).</p></div>',
        unsafe_allow_html=True,
    )

    with st.form("form_login"):
        role = st.radio("Masuk sebagai", ROLES, horizontal=True, key="login_role")
        nama = st.text_input("Nama Lengkap", placeholder="Contoh: Ibu Siti Nurhaliza, S.Pd.")
        kelas = None
        if role == "Guru Kelas":
            kelas = st.selectbox("Kelas yang Diampu", list(range(1, 7)), format_func=lambda i: f"Kelas {i}")
        submitted = st.form_submit_button("Masuk →", use_container_width=True)

    if submitted:
        if not nama.strip():
            st.error("Nama tidak boleh kosong.")
        else:
            st.session_state.auth = {"role": role, "name": nama.strip(), "kelas": kelas}
            if role == "Guru Kelas" and kelas:
                st.session_state.kelas_terpilih = kelas
            st.session_state.page = "landing"
            st.rerun()


def _goto(page_key: str):
    st.session_state.page = page_key
    st.rerun()


# ----------------------------------------------------------------------------
# Menu navigasi utama (persisten di semua halaman, difilter sesuai peran)
#
# Catatan teknis: kunci widget radio disusun dinamis dari halaman aktif
# (f"nav_radio_{page}") sehingga setiap kali `st.session_state.page` berubah
# lewat tombol lain (mis. CTA di Beranda), widget dibuat ulang dari nol dan
# otomatis mengikuti `index` yang baru — tanpa perlu (dan tanpa boleh, sesuai
# aturan Streamlit) menimpa nilai widget yang sudah terlanjur dibuat di run
# yang sama.
# ----------------------------------------------------------------------------
def render_top_nav():
    auth = st.session_state.auth
    visible_pages = dict(NAV_PAGES)
    if not auth or auth["role"] != "Kepala Sekolah":
        visible_pages.pop("👨‍💼 Portal Supervisi", None)

    labels = list(visible_pages.keys())
    key_to_label = {v: k for k, v in visible_pages.items()}
    current_label = key_to_label.get(st.session_state.page, labels[0])
    if current_label not in labels:
        current_label = labels[0]

    col_nav, col_user = st.columns([5, 2])
    with col_nav:
        with st.container(key="top-nav"):
            choice = st.radio(
                "Menu navigasi",
                labels,
                index=labels.index(current_label),
                horizontal=True,
                label_visibility="collapsed",
                key=f"nav_radio_{st.session_state.page}",
            )
    with col_user:
        role_icon = "👨‍💼" if auth["role"] == "Kepala Sekolah" else "🧑‍🏫"
        st.markdown(
            f'<div class="user-chip">{role_icon} {auth["name"].split()[0]}</div>',
            unsafe_allow_html=True,
        )
        if st.button("Keluar", key="btn_logout", use_container_width=True):
            st.session_state.auth = None
            st.session_state.page = "landing"
            st.rerun()

    new_page = visible_pages[choice]
    if new_page != st.session_state.page:
        st.session_state.page = new_page
        st.rerun()


# ----------------------------------------------------------------------------
# Halaman: Landing Page
# ----------------------------------------------------------------------------
def render_landing():
    st.markdown(
        f"""
        <div class="hero-banner">
            <div class="hero-eyebrow">🧪 LMS STEM • Gugus Lengkongjaya</div>
            <div class="hero-title">{OVERVIEW['judul']}</div>
            <div class="hero-subtitle">{OVERVIEW['subjudul']}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="section-title"><span class="icon">📘</span> Gambaran Umum Pelatihan</div>', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="card">
            <p>{OVERVIEW['definisi']}</p>
            <p>{OVERVIEW['pendekatan']}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    goals_html = "".join(
        f'<div class="goal-item"><span class="goal-bullet">{i+1}</span><span>{g}</span></div>'
        for i, g in enumerate(OVERVIEW["tujuan"])
    )
    st.markdown(f'<div class="card">{goals_html}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="footer-note">Sumber: {OVERVIEW["referensi"]}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-title"><span class="icon">📍</span> Kegiatan Hari Ini</div>', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="card">
            <div class="info-grid">
                <div class="label">Tema</div><div class="value">{EVENT['tema']}</div>
                <div class="label">Waktu</div><div class="value">{EVENT['waktu']}</div>
                <div class="label">Sasaran</div><div class="value">{EVENT['sasaran']}</div>
                <div class="label">Kelompok</div><div class="value">{EVENT['kelompok']}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="section-title"><span class="icon">🕒</span> Lini Masa Acara Hari Ini</div>', unsafe_allow_html=True)
    for item in AGENDA:
        color = PRINSIP_WARNA.get(item["prinsip"], "#9CA3AF")
        badge = (
            f'<span class="badge" style="background:{color}">{item["prinsip"]}</span>'
            if item["prinsip"] != "—" else ""
        )
        st.markdown(
            f"""
            <div class="timeline-item" style="border-left-color:{color}">
                <div class="timeline-time">{item['waktu']}</div>
                <div class="timeline-activity">{item['kegiatan']}</div>
                <div class="timeline-meta">{item['bentuk']} • PJ: {item['penanggung_jawab']}</div>
                {badge}
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown(f'<div class="footer-note">Sumber: {EVENT["referensi"]}</div>', unsafe_allow_html=True)

    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
    if st.button("Lanjut ke Modul Kelas →", use_container_width=True, key="btn_lanjut_modul"):
        _goto("modul_kelas")


# ----------------------------------------------------------------------------
# Util bersama: pembacaan berkas .docx
# ----------------------------------------------------------------------------
def _cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def _docx_mtime(filename: str) -> float:
    path = DOCS_DIR / filename
    return path.stat().st_mtime if path.exists() else 0


def _get_identitas_dict(document) -> dict:
    """Ambil tabel Identitas Modul (tabel pertama) sebagai dict label->nilai."""
    if not document.tables:
        return {}
    t = document.tables[0]
    return {
        _cell_text(r.cells[0]): _cell_text(r.cells[1])
        for r in t.rows
        if len(r.cells) >= 2 and _cell_text(r.cells[0])
    }


_HARI_ID = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
_BULAN_ID = [
    "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember",
]


def _format_tanggal_id(d) -> str:
    return f"{_HARI_ID[d.weekday()]}, {d.day} {_BULAN_ID[d.month - 1]} {d.year}"


def _project_dates(n: int, weekday: int = 2, start_after=None):
    """Proyeksikan `n` tanggal mingguan pada hari tertentu (default Rabu=2),
    dipakai sebagai perkiraan jadwal untuk modul yang belum mencantumkan
    tanggal eksplisit pada dokumen sumber (mis. Rumah Cerdas Kelas 6)."""
    import datetime

    base = start_after or datetime.date.today()
    days_ahead = (weekday - base.weekday()) % 7
    first = base + datetime.timedelta(days=days_ahead)
    return [first + datetime.timedelta(weeks=i) for i in range(n)]


_PERT_RE = re.compile(r"^Pertemuan\s*(\d+)\s*–\s*(.+)$")
_DETAIL_START_RE = re.compile(r"^\d+\.\s*Detail Pelaksanaan")
_DETAIL_END_RE = re.compile(r"^\d+\.\s*(Lembar Kerja|LKPD)")
_TAHAP_RE = re.compile(r"^Tahap Praktik Saintifik dan Enjinering\s*\([^)]*\):\s*(.+)$")
_PENGALAMAN_RE = re.compile(
    r"^Pengalaman Belajar:\s*(.+?)\s*\|\s*Prinsip Pembelajaran Mendalam:\s*(.+)$"
)
_JUDUL_TGL_RE = re.compile(
    r"^(?P<tgl>[A-Za-z]+,\s*\d{1,2}\s+\w+\s+\d{4})\s*:\s*(?P<judul>.+?)\s*\((?P<jp>[^)]+)\)\s*$"
)
_JUDUL_TANPA_TGL_RE = re.compile(r"^(?P<judul>.+?)\s*\((?P<jp>[^)]+)\)\s*$")
_SECTION_MAP = [
    (re.compile(r"^Tujuan Spesifik Pertemuan"), "tujuan"),
    (re.compile(r"^Langkah Kegiatan Guru"), "langkah_guru"),
    (re.compile(r"^Langkah Kegiatan Siswa"), "langkah_siswa"),
    (re.compile(r"^Pertanyaan Pemantik"), "pertanyaan"),
    (re.compile(r"^Alat dan Bahan"), "alat_bahan"),
]


def _parse_pertemuan_detail(document):
    """Baca bagian 'Detail Pelaksanaan' pada Modul Proyek STEM dan pecah
    menjadi struktur per pertemuan: tanggal, tahap praktik pedagogis
    (Praktik Saintifik & Enjinering), pengalaman belajar (Siklus 3M),
    prinsip pembelajaran mendalam, tujuan spesifik, langkah kegiatan guru
    & siswa, pertanyaan pemantik, dan alat & bahan — mengikuti urutan asli
    dokumen sehingga tahan terhadap variasi kecil antar modul."""
    pertemuan_list = []
    in_detail = False
    current = None
    current_section = None

    for item in _iter_block_items(document):
        if item.__class__.__name__ != "Paragraph":
            continue
        txt = item.text.strip()
        if not txt:
            continue

        if _DETAIL_START_RE.match(txt):
            in_detail = True
            continue
        if not in_detail:
            continue
        if _DETAIL_END_RE.match(txt):
            break

        m = _PERT_RE.match(txt)
        if m:
            if current:
                pertemuan_list.append(current)
            rest = m.group(2).strip()
            mm = _JUDUL_TGL_RE.match(rest)
            if mm:
                hari_tanggal, judul, jp = mm.group("tgl"), mm.group("judul").strip(), mm.group("jp").strip()
            else:
                mm2 = _JUDUL_TANPA_TGL_RE.match(rest)
                hari_tanggal = None
                judul = mm2.group("judul").strip() if mm2 else rest
                jp = mm2.group("jp").strip() if mm2 else ""
            current = {
                "nomor": int(m.group(1)),
                "hari_tanggal": hari_tanggal,
                "is_perkiraan": False,
                "judul": judul,
                "jp": jp,
                "tahap": "",
                "pengalaman": [],
                "prinsip": [],
                "tujuan": [],
                "langkah_guru": [],
                "langkah_siswa": [],
                "pertanyaan": [],
                "alat_bahan": [],
            }
            current_section = None
            continue

        if current is None:
            continue

        mt = _TAHAP_RE.match(txt)
        if mt:
            current["tahap"] = mt.group(1).strip()
            current_section = None
            continue

        mp = _PENGALAMAN_RE.match(txt)
        if mp:
            current["pengalaman"] = [s.strip() for s in mp.group(1).split(",")]
            current["prinsip"] = [s.strip() for s in mp.group(2).split(",")]
            current_section = None
            continue

        matched_section = False
        for pat, key in _SECTION_MAP:
            if pat.match(txt):
                current_section = key
                matched_section = True
                break
        if matched_section:
            continue

        style_name = item.style.name if item.style else ""
        if current_section and style_name == "List Paragraph":
            current[current_section].append(txt)

    if current:
        pertemuan_list.append(current)

    return pertemuan_list


def _iter_block_items(document):
    """Iterasi paragraf & tabel sesuai urutan aslinya di dalam dokumen Word."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _slug(text: str) -> str:
    return re.sub(r"\W+", "_", text.lower()).strip("_")[:40]


@st.cache_resource(show_spinner="Membuka dokumen...")
def load_docx_document(filename: str, _mtime: float):
    """Buka dokumen .docx mentah (di-cache sebagai resource, bukan data,
    karena objek python-docx tidak cocok di-pickle oleh cache_data)."""
    path = DOCS_DIR / filename
    if not path.exists():
        return None, f"Berkas '{filename}' tidak ditemukan di folder Cowork ({DOCS_DIR})."
    try:
        return docx.Document(str(path)), None
    except Exception as exc:
        return None, f"Gagal membaca '{filename}': {exc}"


def _table_is_blank(table) -> bool:
    """Tabel lembar kerja peserta didik (kosong, untuk diisi tangan saat
    kegiatan) tidak informatif untuk ditampilkan di aplikasi — dilewati."""
    for row in table.rows:
        for cell in row.cells:
            if _cell_text(cell).strip():
                return False
    return True


def render_docx_generic(document, key_prefix: str = "doc"):
    """Tampilkan seluruh isi dokumen .docx (Modul Proyek / Panduan Guru) apa
    adanya, mengikuti urutan aslinya:
    - Paragraf pendek/tebal/berpola "1. Judul Bagian" -> sub-judul.
    - Paragraf biasa -> teks.
    - Tabel 1 sel -> kotak tips/catatan.
    - Tabel kosong (lembar kerja peserta didik) -> dilewati + catatan singkat.
    - Tabel lain -> daftar label:isi (2 kolom) atau tabel data (>2 kolom).
    """
    idx = 0
    for item in _iter_block_items(document):
        idx += 1
        if item.__class__.__name__ == "Paragraph":
            txt = item.text.strip()
            if not txt:
                continue
            is_bold = bool(item.runs) and all(
                (r.bold or not r.text.strip()) for r in item.runs
            )
            heading_like = bool(re.match(r"^(\d+\.\s|[A-Z]\.\s)", txt)) or (is_bold and len(txt) <= 90)
            if heading_like:
                st.markdown(f'<div class="doc-heading">{txt}</div>', unsafe_allow_html=True)
            elif is_bold:
                st.markdown(f'<div class="doc-title">{txt}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="doc-text">{txt}</div>', unsafe_allow_html=True)
            continue

        table = item
        if len(table.rows) == 0:
            continue
        ncols = len(table.columns)

        if ncols == 1 and len(table.rows) <= 2:
            text = " ".join(_cell_text(c) for r in table.rows for c in r.cells).strip()
            if text:
                st.markdown(f'<div class="tip-box">💡 {text}</div>', unsafe_allow_html=True)
            continue

        if _table_is_blank(table):
            st.markdown(
                '<div class="worksheet-note">📄 Bagian ini adalah lembar kerja peserta didik '
                '(diisi langsung saat kegiatan berlangsung).</div>',
                unsafe_allow_html=True,
            )
            continue

        if ncols == 2:
            rows_html = "".join(
                f'<div class="doc-kv"><div class="doc-k">{_cell_text(r.cells[0])}</div>'
                f'<div class="doc-v">{_cell_text(r.cells[1])}</div></div>'
                for r in table.rows
                if _cell_text(r.cells[0]) or _cell_text(r.cells[1])
            )
            if rows_html:
                st.markdown(f'<div class="doc-table-card">{rows_html}</div>', unsafe_allow_html=True)
        else:
            headers = [_cell_text(c) for c in table.rows[0].cells]
            data_rows = [
                dict(zip(headers, [_cell_text(c) for c in r.cells]))
                for r in table.rows[1:]
            ]
            data_rows = [r for r in data_rows if any(v for v in r.values())]
            if data_rows:
                st.dataframe(data_rows, use_container_width=True, hide_index=True, key=f"{key_prefix}_tbl_{idx}")


# ----------------------------------------------------------------------------
# Pembacaan dinamis berkas "Daftar_Proyek_STEM_Kelas*.docx"
# ----------------------------------------------------------------------------
@st.cache_data(show_spinner="Membaca daftar proyek dari dokumen...")
def load_class_projects(kelas: int, _mtime: float):
    filename = CLASS_FILES.get(kelas)
    if not filename:
        return None, f"Tidak ada berkas terdaftar untuk Kelas {kelas}."

    path = DOCS_DIR / filename
    if not path.exists():
        return None, f"Berkas '{filename}' tidak ditemukan di folder Cowork ({DOCS_DIR})."

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        return None, f"Gagal membaca '{filename}': {exc}"

    if not document.tables:
        return None, f"Tidak ditemukan tabel proyek di dalam '{filename}'."

    table = document.tables[0]
    headers = [_cell_text(c) for c in table.rows[0].cells]

    projects = []
    for row in table.rows[1:]:
        values = [_cell_text(c) for c in row.cells]
        if len(values) < len(headers) or not values[2].strip():
            continue
        projects.append(dict(zip(headers, values)))

    return {"filename": filename, "headers": headers, "projects": projects}, None


def _pert_section(title: str, icon: str, items: list):
    if not items:
        return
    st.markdown(f'<div class="pert-section-label">{icon} {title}</div>', unsafe_allow_html=True)
    li = "".join(f"<li>{x}</li>" for x in items)
    st.markdown(f'<ul class="pert-list">{li}</ul>', unsafe_allow_html=True)


def render_pertemuan_detail(pert: dict, total: int):
    """Tampilkan uraian lengkap satu pertemuan: tanggal, tahap praktik
    pedagogis, penanda pengalaman belajar & prinsip pembelajaran mendalam,
    lalu langkah kegiatan guru/siswa, pertanyaan pemantik, dan alat & bahan."""
    tgl = pert["hari_tanggal"] or "Tanggal belum ditetapkan"
    perkiraan_tag = ' <span class="pert-perkiraan-tag">📌 perkiraan</span>' if pert.get("is_perkiraan") else ""

    st.markdown(
        f"""
        <div class="pert-detail-card">
            <div class="pert-detail-header">🗓️ Pertemuan {pert['nomor']} dari {total} — {tgl}{perkiraan_tag}</div>
            <div class="pert-detail-title">{pert['judul']}</div>
            {f'<div class="pert-detail-jp">⏱️ {pert["jp"]}</div>' if pert.get("jp") else ''}
        </div>
        """,
        unsafe_allow_html=True,
    )

    badges = []
    if pert.get("tahap"):
        badges.append(f'<span class="pert-badge pert-badge-tahap">🧭 Praktik Pedagogis: {pert["tahap"]}</span>')
    for p in pert.get("pengalaman", []):
        color = PENGALAMAN_WARNA.get(p, "#0EA5E9")
        badges.append(
            f'<span class="pert-badge" style="background:{color}1F;color:{color};border-color:{color}66">'
            f'💡 Pengalaman Belajar: {p}</span>'
        )
    for p in pert.get("prinsip", []):
        color = PRINSIP_WARNA.get(p, "#9CA3AF")
        badges.append(
            f'<span class="pert-badge" style="background:{color}1F;color:{color};border-color:{color}66">'
            f'🌱 Prinsip: {p}</span>'
        )
    if badges:
        st.markdown(f'<div class="pert-badge-row">{"".join(badges)}</div>', unsafe_allow_html=True)

    _pert_section("Tujuan Spesifik Pertemuan", "🎯", pert.get("tujuan", []))
    _pert_section("Langkah Kegiatan Guru", "👩‍🏫", pert.get("langkah_guru", []))
    _pert_section("Langkah Kegiatan Siswa — Pengalaman Belajar", "🧒", pert.get("langkah_siswa", []))
    _pert_section("Pertanyaan Pemantik", "❓", pert.get("pertanyaan", []))
    _pert_section("Alat dan Bahan", "🧰", pert.get("alat_bahan", []))


def render_project_card(headers, proj: dict, card_key: str = ""):
    values = list(proj.values())
    smt = values[1] if len(values) > 1 else "-"
    judul = values[2] if len(values) > 2 else "(tanpa judul)"
    produk = values[7] if len(values) > 7 else ""

    st.markdown(
        f"""
        <div class="project-card">
            <span class="project-chip">Semester {smt}</span>
            <div class="project-title">{judul}</div>
            {f'<div class="project-produk">🎯 <b>Produk akhir:</b> {produk}</div>' if produk else ''}
        </div>
        """,
        unsafe_allow_html=True,
    )

    # --- Rencana pelaksanaan per pertemuan (jika Modul Proyek lengkap tersedia) ---
    docs = PROJECT_DOCS.get(judul.strip())
    pertemuan_list = []
    identitas = {}
    if docs and docs.get("modul") and docx is not None:
        document, err = load_docx_document(docs["modul"], _docx_mtime(docs["modul"]))
        if not err:
            identitas = _get_identitas_dict(document)
            pertemuan_list = _parse_pertemuan_detail(document)
            if pertemuan_list and not pertemuan_list[0]["hari_tanggal"]:
                # Modul belum mencantumkan tanggal eksplisit (mis. Rumah Cerdas
                # Kelas 6) — lengkapi dengan proyeksi jadwal mingguan agar
                # tanggal pertemuan tetap tuntas untuk seluruh sesi.
                proyeksi = _project_dates(len(pertemuan_list), weekday=2)
                for p, tgl in zip(pertemuan_list, proyeksi):
                    p["hari_tanggal"] = _format_tanggal_id(tgl)
                    p["is_perkiraan"] = True

    if pertemuan_list:
        chips = "".join(
            f'<div class="pertemuan-chip">🗓️ <b>Pertemuan {p["nomor"]}:</b> {p["hari_tanggal"]}'
            f'{" 📌" if p.get("is_perkiraan") else ""}</div>'
            for p in pertemuan_list
        )
        st.markdown(f'<div class="pertemuan-row">{chips}</div>', unsafe_allow_html=True)
        if any(p.get("is_perkiraan") for p in pertemuan_list):
            st.markdown(
                '<div class="pertemuan-note">📌 Tanggal pada modul ini belum ditetapkan eksplisit — '
                'ditampilkan perkiraan jadwal mingguan; sesuaikan dengan Kalender Pendidikan sekolah.</div>',
                unsafe_allow_html=True,
            )

        with st.expander(f"📖 Uraian Langkah Tiap Pertemuan ({len(pertemuan_list)} pertemuan)"):
            sel_key = f"pert_sel_{card_key}"
            if sel_key not in st.session_state:
                st.session_state[sel_key] = pertemuan_list[0]["nomor"]

            with st.container(key=f"pert_menu_{card_key}"):
                cols = st.columns(len(pertemuan_list), gap="small")
                for i, col in enumerate(cols):
                    p = pertemuan_list[i]
                    with col:
                        selected = st.session_state[sel_key] == p["nomor"]
                        if st.button(
                            f"P{p['nomor']}",
                            key=f"btn_pert_{card_key}_{p['nomor']}",
                            use_container_width=True,
                            type="primary" if selected else "secondary",
                        ):
                            st.session_state[sel_key] = p["nomor"]
                            st.rerun()

            chosen = next(
                (p for p in pertemuan_list if p["nomor"] == st.session_state[sel_key]),
                pertemuan_list[0],
            )
            render_pertemuan_detail(chosen, len(pertemuan_list))
    elif identitas.get("Alokasi Waktu"):
        st.markdown(
            f'<div class="pertemuan-note">🗓️ Alokasi waktu: {identitas["Alokasi Waktu"]} '
            f'— tanggal spesifik tiap pertemuan belum dicantumkan pada modul ini.</div>',
            unsafe_allow_html=True,
        )
    else:
        # Proyek ini belum memiliki dokumen Modul Proyek STEM lengkap (baru 1
        # proyek per kelas yang sudah disusun rinci). Tetap tampilkan
        # perkiraan jadwal umum (3 pertemuan mingguan) agar setiap proyek di
        # Daftar Proyek punya gambaran waktu pelaksanaan, dengan penanda
        # jelas bahwa ini perkiraan generik — bukan jadwal resmi dari modul,
        # karena rincian pertemuan/kegiatannya memang belum tersusun.
        _generic_weekday = hash(judul.strip()) % 5  # variasi Senin(0)-Jumat(4) antarproyek
        proyeksi = _project_dates(3, weekday=_generic_weekday)
        chips = "".join(
            f'<div class="pertemuan-chip pertemuan-chip-generic">🗓️ <b>Pertemuan {i+1}:</b> '
            f'{_format_tanggal_id(tgl)} 📌</div>'
            for i, tgl in enumerate(proyeksi)
        )
        st.markdown(f'<div class="pertemuan-row">{chips}</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="pertemuan-note">📌 Perkiraan jadwal umum (3 pertemuan mingguan) — modul proyek '
            'lengkap untuk topik ini belum disusun sehingga rincian langkah kegiatan & tanggal resmi belum '
            'tersedia. Sesuaikan dengan Kalender Pendidikan sekolah bila proyek ini akan digunakan.</div>',
            unsafe_allow_html=True,
        )

    with st.expander("Lihat rincian CP & integrasi mata pelajaran"):
        for idx in (3, 4, 5, 6):
            if idx < len(headers) and idx < len(values) and values[idx]:
                st.markdown(f"**{headers[idx]}**")
                st.caption(values[idx])

    # --- Modul Proyek lengkap & Panduan Guru (jika tersedia untuk proyek ini) ---
    if docs and docx is not None:
        if docs.get("modul"):
            with st.expander("📘 Modul Proyek STEM Lengkap"):
                document, err = load_docx_document(docs["modul"], _docx_mtime(docs["modul"]))
                if err:
                    st.error(err)
                else:
                    render_docx_generic(document, key_prefix=f"{card_key}_modul")
                    st.caption(f"Sumber: {docs['modul']}")
        if docs.get("panduan"):
            with st.expander("🛠️ Panduan Guru: Pembuatan Produk"):
                document, err = load_docx_document(docs["panduan"], _docx_mtime(docs["panduan"]))
                if err:
                    st.error(err)
                else:
                    render_docx_generic(document, key_prefix=f"{card_key}_panduan")
                    st.caption(f"Sumber: {docs['panduan']}")


# ----------------------------------------------------------------------------
# Halaman: Modul Kelas
# ----------------------------------------------------------------------------
def render_modul_kelas():
    auth = st.session_state.auth
    is_guru = bool(auth) and auth["role"] == "Guru Kelas"
    guru_kelas = auth.get("kelas") if is_guru else None

    st.markdown('<div class="section-title"><span class="icon">🏫</span> Modul Kelas</div>', unsafe_allow_html=True)

    if is_guru and guru_kelas:
        # Guru Kelas hanya melihat proyek STEM kelas yang diampunya sendiri —
        # menu pilih kelas 1-6 disembunyikan, tidak ada opsi berpindah kelas.
        st.markdown(
            f'<div class="card"><p>Menampilkan Daftar Proyek STEM untuk kelas yang Anda ampu: '
            f'<b>Kelas {guru_kelas}</b>.</p></div>',
            unsafe_allow_html=True,
        )
        st.session_state.kelas_terpilih = guru_kelas
    else:
        st.markdown(
            '<div class="card"><p>Pilih kelas untuk melihat daftar proyek STEM dan mengunggah laporan praktik.</p></div>',
            unsafe_allow_html=True,
        )
        with st.container(key="kelas-menu"):
            cols = st.columns(6, gap="small")
            for i, col in enumerate(cols, start=1):
                with col:
                    selected = st.session_state.kelas_terpilih == i
                    if st.button(
                        str(i),
                        key=f"btn_kelas_{i}",
                        use_container_width=True,
                        type="primary" if selected else "secondary",
                    ):
                        st.session_state.kelas_terpilih = i
                        st.rerun()

    kelas = st.session_state.kelas_terpilih
    ringkas = KELOMPOK_KELAS[kelas - 1]
    st.markdown(
        f'<div class="kelas-chip">📚 Kelas {kelas} dipilih — {ringkas["modul"]} ({ringkas["fase"]})</div>',
        unsafe_allow_html=True,
    )

    storybook_url = STORYBOOK_LINKS.get(kelas)
    if storybook_url:
        st.markdown(
            f"""
            <div class="storybook-card">
                <div class="storybook-icon">📖</div>
                <div class="storybook-text">
                    <div class="storybook-title">Buku Cerita Kelas {kelas}</div>
                    <div class="storybook-sub">Versi storybook proyek STEM di atas — cocok dibacakan untuk peserta didik</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.link_button("📖 Baca Buku Cerita (Storybook) →", storybook_url, use_container_width=True)

    st.markdown(
        f'<div class="section-title"><span class="icon">📋</span> Daftar Proyek STEM Kelas {kelas}</div>',
        unsafe_allow_html=True,
    )

    if docx is None:
        st.error(
            "Modul `python-docx` belum terpasang. Jalankan `pip install -r requirements.txt`, "
            "lalu muat ulang halaman."
        )
    else:
        filename = CLASS_FILES.get(kelas)
        path = DOCS_DIR / filename if filename else None
        mtime = path.stat().st_mtime if path and path.exists() else 0
        data, err = load_class_projects(kelas, mtime)

        if err:
            st.error(err)
        else:
            st.caption(f"Sumber: {data['filename']} • {len(data['projects'])} proyek ditemukan")
            for i, proj in enumerate(data["projects"]):
                render_project_card(data["headers"], proj, card_key=f"k{kelas}_p{i}")

    st.markdown('<div class="section-title"><span class="icon">📤</span> Unggah Dokumen Guru</div>', unsafe_allow_html=True)
    uploaded = st.file_uploader(
        f"📤 Unggah Laporan Praktik / Foto Produk STEM Kelas {kelas}",
        type=["png", "jpg", "pdf", "docx"],
        key=f"uploader_kelas_{kelas}",
    )
    if uploaded is not None:
        st.success(f"✅ Dokumen '{uploaded.name}' berhasil diunggah ke sistem LMS (simulasi).")


# ----------------------------------------------------------------------------
# Pembacaan lengkap "Program_Supervisi_Kepala_Sekolah_STEM.docx"
# (Lampiran 1-4, mengikuti struktur asli dokumen)
# ----------------------------------------------------------------------------
@st.cache_data(show_spinner="Membaca program supervisi dari dokumen...")
def load_supervisi_full(_mtime: float):
    path = DOCS_DIR / SUPERVISI_FILE
    if not path.exists():
        return None, f"Berkas '{SUPERVISI_FILE}' tidak ditemukan di folder Cowork ({DOCS_DIR})."

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        return None, f"Gagal membaca '{SUPERVISI_FILE}': {exc}"

    tujuan = []
    capture_tujuan = False
    current_lamp = None
    last_heading = None
    capturing_kesimpulan = False
    kesimpulan_buffer = []

    state = {
        1: {"identitas": [], "skala": "", "kategori": [], "catatan_aspek": []},
        2: {"identitas": [], "checklist": [], "kesimpulan_opsi": [], "catatan_aspek": []},
        3: {"identitas": [], "aspek": []},
        4: {"rekap_header": [], "rekap": []},
    }

    for item in _iter_block_items(document):
        cls_name = item.__class__.__name__
        if cls_name == "Paragraph":
            txt = item.text.strip()
            if not txt:
                continue
            if txt.startswith("D. TUJUAN SUPERVISI"):
                capture_tujuan = True
                continue
            if txt.startswith("E. SASARAN SUPERVISI"):
                capture_tujuan = False
            if capture_tujuan and not txt.rstrip().endswith("bertujuan untuk:"):
                tujuan.append(txt)

            m = re.match(r"^Lampiran (\d)\.", txt)
            if m:
                current_lamp = int(m.group(1))
                last_heading = None
                continue

            if txt == "Kesimpulan Telaah":
                capturing_kesimpulan = True
                kesimpulan_buffer = []
                continue
            if capturing_kesimpulan:
                if re.match(r"^[A-E]\.\s", txt) or txt.startswith("Lampiran"):
                    state[2]["kesimpulan_opsi"] = kesimpulan_buffer
                    capturing_kesimpulan = False
                else:
                    kesimpulan_buffer.append(txt)
                    continue

            if current_lamp == 1 and re.match(r"^[A-E]\.\s", txt):
                last_heading = txt
        else:  # Table
            headers = [_cell_text(c) for c in item.rows[0].cells]
            if capturing_kesimpulan:
                state[2]["kesimpulan_opsi"] = kesimpulan_buffer
                capturing_kesimpulan = False

            if current_lamp == 1:
                if headers and headers[0].startswith("Skala Penilaian"):
                    state[1]["skala"] = headers[0]
                elif headers[:1] == ["Indikator"]:
                    inds = [_cell_text(r.cells[0]) for r in item.rows[1:] if _cell_text(r.cells[0])]
                    if inds:
                        state[1]["kategori"].append({"kategori": last_heading or "", "indikator": inds})
                elif headers and headers[0] == "Kolom Isian":
                    state[1]["identitas"] = [_cell_text(r.cells[0]) for r in item.rows[1:]]
                elif headers == ["Aspek", "Uraian"]:
                    state[1]["catatan_aspek"] = [_cell_text(r.cells[0]) for r in item.rows[1:]]
            elif current_lamp == 2:
                if headers and headers[0] == "Kolom Isian":
                    state[2]["identitas"] = [_cell_text(r.cells[0]) for r in item.rows[1:]]
                elif len(headers) > 1 and headers[0] == "No" and "Komponen" in headers[1]:
                    for r in item.rows[1:]:
                        state[2]["checklist"].append([_cell_text(c) for c in r.cells])
                elif headers == ["Aspek", "Uraian"]:
                    state[2]["catatan_aspek"] = [_cell_text(r.cells[0]) for r in item.rows[1:]]
            elif current_lamp == 3:
                if headers and headers[0] == "Kolom Isian":
                    state[3]["identitas"] = [_cell_text(r.cells[0]) for r in item.rows[1:]]
                elif headers == ["Aspek", "Uraian"]:
                    state[3]["aspek"] = [_cell_text(r.cells[0]) for r in item.rows[1:]]
            elif current_lamp == 4:
                if headers and headers[0] == "No" and "Kelas" in headers:
                    state[4]["rekap_header"] = headers
                    for r in item.rows[1:]:
                        state[4]["rekap"].append([_cell_text(c) for c in r.cells])

    if not state[1]["kategori"]:
        return None, f"Tidak ditemukan instrumen Lampiran 1 di dalam '{SUPERVISI_FILE}'."

    return {"tujuan": tujuan, "lampiran": state}, None


# ----------------------------------------------------------------------------
# Widget dinamis untuk field identitas (label diambil langsung dari dokumen)
# ----------------------------------------------------------------------------
def render_identity_field(label: str, key_prefix: str, prefill: str = ""):
    low = label.lower()
    key = f"{key_prefix}_{_slug(label)}"

    if "kelas" in low and "semester" in low:
        c1, c2 = st.columns(2)
        with c1:
            kelas_val = st.selectbox("Kelas", [f"Kelas {i}" for i in range(1, 7)], key=f"{key}_kelas")
        with c2:
            smt_val = st.selectbox("Semester", ["Ganjil", "Genap"], key=f"{key}_smt")
        return f"{kelas_val} / {smt_val}"
    if "tanggal" in low or "hari, tanggal" in low:
        return str(st.date_input(label, key=key))
    if "tahap 3m" in low or "pertemuan ke-" in low:
        return st.selectbox(label, ["Memahami", "Mengaplikasi", "Merefleksi"], key=key)
    if "jenis supervisi" in low:
        return st.selectbox(label, ["Dokumen", "Kelas/Akademik"], key=key)
    return st.text_input(label, value=prefill, key=key)


# ----------------------------------------------------------------------------
# Lampiran 1 — Instrumen Observasi Kelas STEM PjBL
# ----------------------------------------------------------------------------
def render_lampiran1(data):
    lamp = data["lampiran"][1]
    auth = st.session_state.auth

    st.markdown('<div class="section-title"><span class="icon">📋</span> Identitas Observasi</div>', unsafe_allow_html=True)
    with st.expander("Isi identitas observasi", expanded=False):
        for label in lamp["identitas"]:
            prefill = auth["name"] if "kepala sekolah" in label.lower() else ""
            render_identity_field(label, "l1", prefill)

    if lamp["skala"]:
        st.markdown(f'<div class="skala-note">📏 {lamp["skala"]}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-title"><span class="icon">✍️</span> Penilaian per Indikator</div>', unsafe_allow_html=True)
    skor_total = []
    for kat in lamp["kategori"]:
        with st.expander(kat["kategori"], expanded=False):
            for i, ind in enumerate(kat["indikator"]):
                st.markdown(f'<div class="indikator-text">{ind}</div>', unsafe_allow_html=True)
                skor = st.select_slider(
                    "Skor", options=[1, 2, 3, 4], value=3,
                    key=f"l1_skor_{_slug(kat['kategori'])}_{i}", label_visibility="collapsed",
                )
                skor_total.append(skor)

    st.markdown('<div class="section-title"><span class="icon">📝</span> Catatan Hasil Observasi</div>', unsafe_allow_html=True)
    catatan_vals = {}
    for label in lamp["catatan_aspek"]:
        catatan_vals[label] = st.text_area(label, key=f"l1_catatan_{_slug(label)}")

    if st.button("💾 Simpan Instrumen Observasi Kelas", use_container_width=True, key="btn_simpan_l1"):
        rata2 = sum(skor_total) / len(skor_total) if skor_total else 0
        st.success(f"✅ Instrumen Observasi Kelas berhasil disimpan! Rata-rata skor: **{rata2:.2f} / 4**.")
        st.balloons()


# ----------------------------------------------------------------------------
# Lampiran 2 — Checklist Telaah Dokumen Modul Proyek STEM (10 Komponen)
# ----------------------------------------------------------------------------
def render_lampiran2(data):
    lamp = data["lampiran"][2]
    auth = st.session_state.auth

    st.markdown('<div class="section-title"><span class="icon">📋</span> Identitas Telaah</div>', unsafe_allow_html=True)
    with st.expander("Isi identitas telaah", expanded=False):
        for label in lamp["identitas"]:
            prefill = auth["name"] if "penelaah" in label.lower() else ""
            render_identity_field(label, "l2", prefill)

    st.markdown('<div class="section-title"><span class="icon">✅</span> Checklist 10 Komponen</div>', unsafe_allow_html=True)
    hasil_checklist = []
    for row in lamp["checklist"]:
        no, komponen = row[0], row[1]
        st.markdown(f'<div class="indikator-text"><b>{no}.</b> {komponen}</div>', unsafe_allow_html=True)
        sesuai = st.radio(
            "Sesuai?", ["Ya", "Tidak"], horizontal=True,
            key=f"l2_sesuai_{no}", label_visibility="collapsed",
        )
        hasil_checklist.append(sesuai)

    if lamp["kesimpulan_opsi"]:
        st.markdown('<div class="section-title"><span class="icon">🏁</span> Kesimpulan Telaah</div>', unsafe_allow_html=True)
        st.selectbox("Kesimpulan", lamp["kesimpulan_opsi"], key="l2_kesimpulan", label_visibility="collapsed")

    st.markdown('<div class="section-title"><span class="icon">📝</span> Catatan Penelaah</div>', unsafe_allow_html=True)
    for label in lamp["catatan_aspek"]:
        st.text_area(label, key=f"l2_catatan_{_slug(label)}")

    if st.button("💾 Simpan Checklist Telaah Dokumen", use_container_width=True, key="btn_simpan_l2"):
        jumlah_ya = hasil_checklist.count("Ya")
        st.success(
            f"✅ Checklist Telaah Dokumen berhasil disimpan! {jumlah_ya}/{len(hasil_checklist)} "
            f"komponen dinilai sesuai."
        )
        st.balloons()


# ----------------------------------------------------------------------------
# Lampiran 3 — Lembar Tindak Lanjut / Coaching Pasca-Supervisi
# ----------------------------------------------------------------------------
def render_lampiran3(data):
    lamp = data["lampiran"][3]

    st.markdown('<div class="section-title"><span class="icon">📋</span> Identitas Supervisi</div>', unsafe_allow_html=True)
    with st.expander("Isi identitas supervisi", expanded=True):
        for label in lamp["identitas"]:
            render_identity_field(label, "l3")

    st.markdown('<div class="section-title"><span class="icon">🤝</span> Tindak Lanjut / Coaching</div>', unsafe_allow_html=True)
    for label in lamp["aspek"]:
        st.text_area(label, key=f"l3_aspek_{_slug(label)}")

    if st.button("💾 Simpan Lembar Tindak Lanjut", use_container_width=True, key="btn_simpan_l3"):
        st.success("✅ Lembar Tindak Lanjut / Coaching berhasil disimpan ke sistem LMS!")
        st.balloons()


# ----------------------------------------------------------------------------
# Lampiran 4 — Rekapitulasi Hasil Supervisi Satu Tahun Ajaran
# ----------------------------------------------------------------------------
def render_lampiran4(data):
    lamp = data["lampiran"][4]
    st.markdown(
        '<div class="card"><p>Rekapitulasi ini bersifat baca-saja, menampilkan jadwal supervisi dokumen '
        'dan akademik satu tahun ajaran sesuai dokumen. Kolom Ringkasan Hasil & Status Tindak Lanjut '
        'diisi manual pada dokumen resmi setelah supervisi dilaksanakan.</p></div>',
        unsafe_allow_html=True,
    )

    if not lamp["rekap"]:
        st.info("Tidak ada data rekapitulasi ditemukan.")
        return

    headers = lamp["rekap_header"]
    rows = [dict(zip(headers, r)) for r in lamp["rekap"]]
    st.dataframe(rows, use_container_width=True, hide_index=True)
    st.caption(f"Total {len(rows)} jadwal supervisi tercatat untuk tahun ajaran ini.")


# ----------------------------------------------------------------------------
# Halaman: Portal Supervisi Kepala Sekolah
# ----------------------------------------------------------------------------
def render_supervisi():
    st.markdown('<div class="section-title"><span class="icon">👨‍💼</span> Portal Supervisi Kepala Sekolah</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="card"><p>Instrumen supervisi lengkap penerapan pembelajaran STEM berbasis proyek Kelas 1–6, '
        'dibaca langsung dari Program Supervisi Kepala Sekolah.</p></div>',
        unsafe_allow_html=True,
    )

    if docx is None:
        st.error(
            "Modul `python-docx` belum terpasang. Jalankan `pip install -r requirements.txt`, "
            "lalu muat ulang halaman."
        )
        return

    path = DOCS_DIR / SUPERVISI_FILE
    mtime = path.stat().st_mtime if path.exists() else 0
    data, err = load_supervisi_full(mtime)

    if err:
        st.error(err)
        return

    if data["tujuan"]:
        with st.expander("🎯 Tujuan Supervisi", expanded=False):
            for i, g in enumerate(data["tujuan"]):
                st.markdown(f"{i+1}. {g}")

    # --- Sub-navigasi Lampiran 1-4 ---
    labels = list(LAMPIRAN_TABS.keys())
    key_to_label = {v: k for k, v in LAMPIRAN_TABS.items()}
    current_label = key_to_label.get(st.session_state.lampiran_tab, labels[0])

    with st.container(key="lampiran-menu"):
        cols = st.columns(4, gap="small")
        for col, label in zip(cols, labels):
            with col:
                selected = current_label == label
                if st.button(
                    label, key=f"btn_tab_{LAMPIRAN_TABS[label]}", use_container_width=True,
                    type="primary" if selected else "secondary",
                ):
                    st.session_state.lampiran_tab = LAMPIRAN_TABS[label]
                    st.rerun()

    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    tab = st.session_state.lampiran_tab
    if tab == "lampiran1":
        render_lampiran1(data)
    elif tab == "lampiran2":
        render_lampiran2(data)
    elif tab == "lampiran3":
        render_lampiran3(data)
    else:
        render_lampiran4(data)

    st.markdown(f'<div class="footer-note">Sumber: {SUPERVISI_FILE}</div>', unsafe_allow_html=True)


# ----------------------------------------------------------------------------
# Router
# ----------------------------------------------------------------------------
def main():
    load_css()
    init_state()

    if not st.session_state.auth:
        render_login()
        return

    render_top_nav()

    if st.session_state.page == "supervisi" and st.session_state.auth["role"] != "Kepala Sekolah":
        st.session_state.page = "landing"

    if st.session_state.page == "landing":
        render_landing()
    elif st.session_state.page == "modul_kelas":
        render_modul_kelas()
    else:
        render_supervisi()


if __name__ == "__main__":
    main()
